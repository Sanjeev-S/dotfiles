#!/usr/bin/env python3
"""Build a Word version of the City Surf Project FY27 fundraising proposal."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips, Cm, Emu

NAVY = RGBColor(0x0B, 0x4A, 0x56)
TEAL = RGBColor(0x1A, 0x6B, 0x7A)
INK = RGBColor(0x1A, 0x24, 0x28)
MUTED = RGBColor(0x5C, 0x6B, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SAND = "F7F1EA"
ROW_ALT = "F4F8F8"
HEADER_BG = "0B4A56"
LEDE_BG = "EEF5F6"


def set_run(run, *, size=10, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), kwargs.get("val", "single"))
        el.set(qn("w:sz"), kwargs.get("sz", "4"))
        el.set(qn("w:color"), kwargs.get("color", "D5DEE0"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para(doc, text="", *, size=10.5, bold=False, color=INK, space_after=6, space_before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_runs(p, parts):
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run(r, **kwargs)


def heading(doc, text):
    p = para(doc, text, size=13, bold=True, color=NAVY, space_before=10, space_after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B4A56")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def subhead(doc, text):
    return para(doc, text, size=11, bold=True, color=TEAL, space_before=6, space_after=2)


def bullet(doc, label, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(label)
    set_run(r, size=10.5, bold=True, color=INK)
    r2 = p.add_run(" " + body)
    set_run(r2, size=10.5, color=INK)
    return p


def plain_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)
    return p


def make_table(doc, headers, rows, col_widths, numeric=None):
    numeric = numeric or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(w)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, HEADER_BG)
        set_cell_border(cell, color="0B4A56", sz="4")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_run(r, size=9, bold=True, color=WHITE)
        if i in numeric:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for r_i, row in enumerate(rows):
        is_total = row[-1] == "TOTAL"
        values = row[:-1] if is_total else row
        # allow last flag
        if len(row) == len(headers) + 1 and row[-1] in ("TOTAL",):
            values = row[:-1]
        else:
            values = row
            is_total = False
        for c_i, val in enumerate(values):
            cell = table.rows[r_i + 1].cells[c_i]
            fill = "EEF5F6" if is_total else (ROW_ALT if r_i % 2 else "FFFFFF")
            shade(cell, fill)
            set_cell_border(cell, color="D5DEE0", sz="4")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            set_run(r, size=9, bold=is_total, color=INK)
            if c_i in numeric:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_narrow_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("City Surf Project  ·  Confidential  ·  Development Director work product  ·  ")
    set_run(r, size=8, color=MUTED)
    # PAGE field
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run(r2, size=8, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    set_narrow_margins(section)
    add_footer(section)

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10.5)
    styles.font.color.rgb = INK

    k = para(doc, "MEMORANDUM", size=9, bold=True, color=NAVY, space_after=2)
    k.runs[0].font.all_caps = True

    para(doc, "From $1.1 million to $1.4 million", size=20, bold=True, color=NAVY, space_after=1)
    para(
        doc,
        "A fundraising plan for City Surf Project’s next operating year",
        size=12,
        color=TEAL,
        space_after=8,
    )

    meta = [
        ("To", "Johnny Irwin, Executive Director, and the Board Development Committee"),
        ("From", "Sanjeev Suresh, candidate for Development Director"),
        ("Date", "August 14, 2026"),
        ("Re", "FY27 revenue strategy — what to grow, what to hold, and the first 90 days"),
    ]
    for label, value in meta:
        p = para(doc, "", size=10.5, space_after=1)
        add_runs(
            p,
            [
                (f"{label}  ", {"size": 10.5, "bold": True, "color": MUTED}),
                (value, {"size": 10.5, "color": INK}),
            ],
        )

    lede = para(doc, "", size=10.5, space_before=8, space_after=8)
    add_runs(
        lede,
        [
            ("The plan in one sentence. ", {"size": 10.5, "bold": True, "color": NAVY}),
            (
                "Grow individual major gifts and corporate cash — not events, not government — and use the first 90 days to stand up a Wave Makers portfolio that can produce half of the $300,000 gap.",
                {"size": 10.5, "color": INK},
            ),
        ],
    )

    heading(doc, "1. Assessment: the work is ready; the revenue mix is not")
    para(
        doc,
        "City Surf Project has earned the right to grow. In ten years you have put more than 3,000 young people in the water, partnered with a dozen SFUSD schools from Mission to Leadership to Ida B. Wells, and been named a 2025 California Nonprofit of the Year. For 80% of students, their first time in the ocean is with CSP. After one outing, 96% feel closer to their peers; 71% come to see themselves as surfers. That is a fundraising story most $1 million organizations would envy.",
    )
    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            (
                "The revenue model has not kept pace. Form 990 for the year ended May 2025 shows revenue of ",
                {"size": 10.5},
            ),
            ("$1,110,667", {"size": 10.5, "bold": True}),
            (" and expenses of ", {"size": 10.5}),
            ("$1,134,574", {"size": 10.5, "bold": True}),
            (
                " — a $23,907 deficit, the latest in a run of deficit years that cut net assets from $723,884 (calendar 2021) to ",
                {"size": 10.5},
            ),
            ("$362,119", {"size": 10.5, "bold": True}),
            (
                ". That is about 3.8 months of reserve on a $1.13 million operating budget. Fundraising consumed roughly 22% of spending in the prior year, high for this budget size, while the organization employs 21 people and already fields four development seats.",
                {"size": 10.5},
            ),
        ],
    )
    para(doc, "Your current mix on a $1.1 million goal is the real diagnostic:", space_after=4)

    make_table(
        doc,
        ["Stream", "Share", "≈ Dollars", "Read"],
        [
            ["Individual donors", "30%", "$330,000", "Largest flexible stream — under-cultivated at $2,500+"],
            ["Foundation / grants", "20%", "$220,000", "Real, but leaky: several six-figure relationships look lapsed"],
            ["Corporate sponsorship", "3%", "$33,000", "The structural hole. Brands are present; operating cash is not."],
            ["Events", "10%", "$110,000", "Three events already. FY25 net fundraising income on the 990: $8."],
            ["Other (gov’t, earned)", "37%", "$407,000", "Pays the bills. Will not fund a 27% increase. Can shrink."],
        ],
        [1.7, 0.7, 1.1, 3.5],
        numeric={1, 2},
    )

    subhead(doc, "Biggest risks")
    bullet(
        doc,
        "Government concentration.",
        "DCYF funded Youth Surfing at $150,000 against a $323,741 request and listed SILT in the same 2024–29 cycle — after the City said it made “tough decisions” in a deficit year. Whale Tail added $50,000. Precious money. Not a growth engine.",
    )
    bullet(
        doc,
        "Corporate is a rounding error.",
        "Three percent — $33,000 — for a photogenic, equity-centered ocean program in San Francisco is the clearest miss. Surf-a-Thon partners (GoPro, Sun Bum, Mollusk, Specialized, Dryrobe) give product and logos, not payroll.",
    )
    bullet(
        doc,
        "Events work hard for modest net.",
        "You already run three: Surf-a-Thon (2026 goal $75,000, raised $45,000), the August 29 dinner, and Back to the Beach on November 6. They fill the calendar and the cost ratio. They do not fill a $300,000 gap.",
    )
    bullet(
        doc,
        "A growth hire on thin reserves.",
        "A Director at $105–115k is the right seat only if the team is pointed at high-leverage work. Four development people on $1.1 million is already a large shop.",
    )
    bullet(
        doc,
        "Lapsed institutions.",
        "Kaiser ($75,000 through 2022), Horace W. Goldsmith ($100,000 through 2023), and Parks California ($81,000 through 2023) drop out of recent filings. The 20% foundation share is real — and leaky.",
    )

    subhead(doc, "Biggest opportunities")
    bullet(
        doc,
        "Individuals are larger than they look.",
        "Roughly 46% of recorded “foundation” dollars flow through DAFs and pass-throughs. Endaoment alone granted $320,000 over three years. That is major-donor money wearing a foundation label. The file has capacity. It does not yet have a program.",
    )
    bullet(
        doc,
        "This board can open doors.",
        "Public listings include wealth management (Laird Norton Wetherby), venture (Offline Ventures), insurance (New York Life), law (Conger; K&L Gates), and tech (Stack AI), plus two directors seated in 2026. That is a give/get board, not a ceremonial one.",
    )
    bullet(
        doc,
        "The case is timely.",
        "Ten-year mark, Nonprofit of the Year, 434 students and 131 surf days last year, a dinner on August 29. Donors fund momentum. You have it.",
    )
    bullet(
        doc,
        "A grants function already exists.",
        "Public staff listings include a Grants Manager. Foundation growth is a retention-and-replacement job, not a new department the Director has to build from zero.",
    )

    heading(doc, "2. What I would bolster — and what I would not")
    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            ("I would put about 80% of new-dollar effort into ", {"size": 10.5}),
            ("individual major gifts", {"size": 10.5, "bold": True}),
            (" and ", {"size": 10.5}),
            ("corporate cash", {"size": 10.5, "bold": True}),
            (
                ". Together they should produce $250,000 of the $300,000. Foundations add $50,000 by replacing what lapsed. Events and government are held, not grown.",
                {"size": 10.5},
            ),
        ],
    )

    make_table(
        doc,
        ["Stream", "Now", "FY27", "Change", "How the dollars actually appear"],
        [
            ["Individuals", "$330k", "$480k", "+$150k", "Wave Makers: upgrades, DAF asks, 12–18 new gifts of $2,500–$25,000"],
            ["Corporate", "$33k", "$130k", "+$97k", "Five cash packages at $15–25k; two entry sponsors at $5–10k"],
            ["Foundations", "$220k", "$270k", "+$50k", "Renew Marin, Newhall, Lux, Bothin, Olympic Club; replace Kaiser / Goldsmith"],
            ["Events", "$110k", "$120k", "+$10k", "Surf-a-Thon back toward $60k; dinner and concert as cultivation, not expansion"],
            ["Other (gov’t / earned)", "$407k", "$400k", "−$7k", "Protect DCYF and city contracts. No growth target."],
            ["Total", "$1.10M", "$1.40M", "+$300k", "A surplus year — first job of $1.4M is to stop drawing down reserves", "TOTAL"],
        ],
        [1.5, 0.75, 0.75, 0.8, 3.2],
        numeric={1, 2, 3},
    )

    para(
        doc,
        "A four-person team cannot run a new gala, a new government RFP, a 40-proposal grant calendar, and a major-gifts program at once. Something has to be the job. Major gifts and corporate close on a 3–9 month cycle; new government and most new foundations do not.",
    )
    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            ("What I would not do in year one: ", {"size": 10.5, "bold": True}),
            (
                "add a fourth event; hire a fifth development FTE; launch a capital campaign; or treat DCYF as the way we get to $1.4 million. Those are tempting, familiar, and wrong for this shop in this year.",
                {"size": 10.5},
            ),
        ],
    )

    heading(doc, "3. Three priorities for the first 6–12 months")
    subhead(doc, "Priority 1 — Stand up Wave Makers (major gifts). This is first on purpose.")
    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            (
                "It is the only strategy that can produce six figures in year one from people who already know you, and it is the one that uses the board. Target: individuals from $330k to $480k. At this budget, ",
                {"size": 10.5},
            ),
            ("$2,500 is a major gift", {"size": 10.5, "bold": True}),
            (" and ", {"size": 10.5}),
            ("$10,000 is a leadership gift", {"size": 10.5, "bold": True}),
            (
                ". A $10,000 gift funds a classroom’s worth of ocean days; $25,000 underwrites a school partnership. Those are asks a donor can see.",
                {"size": 10.5},
            ),
        ],
    )

    subhead(doc, "Priority 2 — Convert corporate from swag to cash.")
    para(
        doc,
        "Three percent is the structural hole, but corporate fiscal years lock in the fall. Closes will cluster in months 4–12, which is why this is second, not first. Sell five named packages — a school partnership, a summer-camp week, a SILT cohort — at $15–25k, plus two $5–10k entry sponsors. Natural prospects: the action-sports brands already around Surf-a-Thon; Kaiser, whose health frame matches the mental-health outcomes and whose last grant was 2022; REI (already $10,000 through the Cooperative Action Fund); and SF tech and finance introductions from the board. Team-building surf days stay in the menu as a door-opener, not the product we are selling.",
    )

    subhead(doc, "Priority 3 — Defend the base so growth is not canceled by leakage.")
    para(
        doc,
        "The Grants Manager owns renewals (Marin Community Foundation, Henry Mayo Newhall, Miranda Lux, Bothin, Olympic Club) and a short replacement list for Kaiser and Goldsmith. I support two or three site visits. I do not spend my year writing 40 proposals. Government: on-time reporting, no surprises, no growth target. If we raise $300,000 in new private money and lose $80,000 in city or foundation renewals, we did not have a good year.",
    )

    heading(doc, "4. A concrete first move: 90-day Wave Makers launch")
    para(
        doc,
        "I am not proposing a new event. I am proposing that we finally treat $2,500+ donors as a portfolio — with names, next asks, and a weekly meeting rhythm. If I start in September, the August 29 dinner is either week-one work or a 72-hour follow-up list. Either way, it is the first cultivation asset, not a one-night fundraiser.",
    )

    subhead(doc, "Days 1–30 — See the file")
    plain_bullet(doc, "Pull every gift of $500+ for five years. Recode DAF and Endaoment gifts as individual prospects. They have been hiding in the foundation column.")
    plain_bullet(doc, "Wealth-screen the list (DonorSearch or iWave, about $2,500).")
    plain_bullet(doc, "Build a working portfolio of 40: 15 current, 15 lapsed or DAF, 10 new via the board.")
    plain_bullet(doc, "One-on-ones with every board member: personal gift, two introductions, which asks they will walk into with me.")
    plain_bullet(doc, "Sit in on programs. The ask is a student at Linda Mar, not a spreadsheet.")

    subhead(doc, "Days 31–60 — Be in the room")
    plain_bullet(doc, "Twenty donor meetings. The Executive Director joins the $10,000+ conversations — one a week, not a roadshow.")
    plain_bullet(doc, "Dinner follow-up is assigned seating in reverse: handwritten notes, a student story, and a specific next ask within 72 hours. Tickets were the invitation. The gift is the point.")
    plain_bullet(doc, "Launch Wave Makers levels: $2,500 / $5,000 / $10,000 / $25,000, with a two-year option so we are not re-asking from zero every June.")
    plain_bullet(doc, "Board: 100% personal gift by day 45. We cannot ask the city for money the board has not put in.")

    subhead(doc, "Days 61–90 — Close, then lock the system")
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_runs(
        p,
        [
            ("Fifteen more meetings. Success looks like ", {"size": 10.5}),
            ("$50,000–$75,000 closed or verbally committed", {"size": 10.5, "bold": True}),
            (" — proof the model works before we spend the winter on corporate.", {"size": 10.5}),
        ],
    )
    plain_bullet(doc, "Stewardship rule, no exceptions: 48-hour thank-you; ED call at $1,000+; student note or 30-second video at $2,500+.")
    plain_bullet(doc, "November 6 concert: the VIP list is the portfolio. We do not use a Grateful Dead night to acquire $40 tickets. We use it to thank people who can give $5,000.")
    plain_bullet(doc, "A written FY27 calendar on the Committee’s desk: monthly portfolio reviews, quarterly board ask reports, January corporate push.")

    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            ("What I need — and what I do not. ", {"size": 10.5, "bold": True}),
            (
                "Fifty percent of my time; the Development Assistant for about ten hours a week on research and scheduling; the Data Entry Specialist to code the portfolio so we stop living in spreadsheets; the ED for four to six joint visits a quarter; the Development Committee monthly. Cash budget: ",
                {"size": 10.5},
            ),
            ("$4,000", {"size": 10.5, "bold": True}),
            (
                " (screening $2,500; packets and coffee $1,000; student thank-you videos $500). I do not need a new CRM in year one, a consultant, another event, or another hire. If the current database cannot produce a $500+ giving history, that is the first operations fix — still cheaper than software we will not use.",
                {"size": 10.5},
            ),
        ],
    )

    heading(doc, "5. How we will know after year one")
    para(
        doc,
        "If this is working, FY27 looks like the table below — not just a $1.4 million top line. If we hit revenue because a one-time government or DAF gift landed, and the mix did not move, I will call that a miss. The point of this plan is an engine that still works if City Hall has a bad year.",
    )

    make_table(
        doc,
        ["Metric", "Year-one target", "Why this one"],
        [
            ["Total revenue", "$1.40 million", "The number you set. A surplus year; reserves start to rebuild."],
            ["Individual revenue", "$480,000", "Primary growth engine. If this misses, the plan missed."],
            ["Corporate cash", "$130,000", "Proof we converted brands into operating support."],
            ["Donors at $2,500+", "25 (from an assumed thin base)", "Portfolio health. Count matters more than one hero gift."],
            ["Donors at $10,000+", "10", "Leadership layer. These gifts fund school partnerships."],
            ["Board give / get", "100% personal gift; 2 intros each", "Culture change. Optional board giving is optional revenue."],
            ["$1,000+ retention", "80%+", "Cheaper than acquisition. Tells us stewardship is real."],
            ["Months of reserve", "4.5+, path to 6", "Deficits stop. The $1.4M is not all for new programming."],
            ["Fundraising cost ratio", "At or below 20% (from ~22%)", "Growth without a more expensive development shop."],
        ],
        [2.1, 2.3, 2.6],
    )

    p = para(doc, "", space_after=8)
    add_runs(
        p,
        [
            ("Assumptions, so we can correct them. ", {"size": 9.5, "bold": True, "color": MUTED}),
            (
                "(1) “This year” is FY26 (June 2025–May 2026) at $1.1M; “next year” is FY27 at $1.4M. (2) I assume a September 2026 start. (3) The $330k individual figure is mostly sub-$1,000 gifts with a thin major-gifts layer — typical at this stage, and the reason a portfolio will move the number. (4) The fourth development seat is a grants function (public materials list a Grants Manager and a Development Manager); grants stay there. (5) Board giving and the current $2,500+ donor count are not in the public record. Day-30 pipeline report to the Committee replaces these guesses with names.",
                {"size": 9.5, "color": MUTED},
            ),
        ],
    )

    para(
        doc,
        "I would rather be measured on whether the mix moved than on whether we found one large check. City Surf Project does not have a story problem. It has a concentration problem and an underbuilt major-gifts muscle. Those are solvable in a year — if we choose them.",
        italic=True,
        space_after=12,
    )

    doc.add_page_break()
    para(doc, "SUPPORTING MATERIALS", size=9, bold=True, color=NAVY, space_after=2)
    para(doc, "Appendix — budget, timeline, and sources", size=18, bold=True, color=NAVY, space_after=1)
    para(
        doc,
        "Not part of the 3-page memo. Use as a working attachment for the Committee.",
        size=11,
        color=TEAL,
        space_after=8,
    )

    heading(doc, "A. Wave Makers — 90-day operating budget")
    make_table(
        doc,
        ["Item", "Owner", "Cost", "Notes"],
        [
            ["Wealth screening (DonorSearch / iWave, 12-month)", "DD + Data Entry", "$2,500", "Screen $500+ donors and board networks, not the whole list"],
            ["Donor meetings (coffee, parking, two small host breakfasts)", "DD", "$1,000", "35 meetings in 90 days. No restaurant cultivation circuit."],
            ["Student thank-you videos and printed packets", "Dev Assistant + program staff", "$500", "Phone videos are enough. Do not hire a filmmaker."],
            ["New CRM, consultant, extra event, extra FTE", "—", "$0", "Explicitly out of scope for the first 90 days"],
            ["Total incremental cash", "", "$4,000", "If $50–75k closes, return is 12–18× in the first quarter", "TOTAL"],
        ],
        [2.6, 1.3, 0.8, 2.3],
        numeric={2},
    )
    para(
        doc,
        "Staff time is the real cost: ~50% of the Director, ~10 hours/week of the Assistant, coding time from Data Entry, and four to six ED visits. That is a reallocation, not a new line.",
        size=10,
        color=MUTED,
    )

    heading(doc, "B. Twelve-month calendar (FY27)")
    make_table(
        doc,
        ["Window", "Major gifts", "Corporate", "Base (grants / gov’t / events)"],
        [
            [
                "Sep–Nov (days 1–90)",
                "Build portfolio of 40. Board 100% in. Close $50–75k. Dinner + concert as cultivation.",
                "Rewrite the partnership deck around cash packages. Board intro list of 15 companies.",
                "Grants Manager: renewal calendar. Surf-a-Thon debrief: why $45k vs. $75k.",
            ],
            [
                "Dec–Feb",
                "Year-end and DAF push. Multi-year Wave Maker asks. ED calls every $1k+ donor.",
                "Proposals into Q1 corporate budgets. Two site visits for priority sponsors.",
                "Foundation LOIs to replace Kaiser / Goldsmith. DCYF reporting on time.",
            ],
            [
                "Mar–May",
                "Spring upgrades. Portfolio to 50. First retention report to Committee.",
                "Close three of five $15–25k packages before fiscal year-end.",
                "Newhall / Marin / Lux renewals. Prep summer-camp surfership messaging.",
            ],
            [
                "Jun–Aug",
                "Stewardship visits at the beach — donors watch a session, then we ask.",
                "Remaining closes. Team-build days only as a door to a cash ask.",
                "Surf-a-Thon run as a $60k campaign with fewer, larger teams — not a broader event.",
            ],
        ],
        [1.4, 2.1, 2.05, 2.15],
    )

    heading(doc, "C. Why these numbers are realistic, not textbook")
    p = para(doc, "", space_after=6)
    add_runs(
        p,
        [
            ("$150,000 in new individual giving is roughly ", {"size": 10.5}),
            ("fifteen gifts that average $10,000", {"size": 10.5, "bold": True}),
            (
                ", or a larger number of $2,500–$5,000 upgrades plus a handful of leadership gifts. On a board that includes wealth management, venture, and insurance — and a file that has already moved six figures through Endaoment — that is a stretch and a plan, not a fantasy. $97,000 in new corporate cash is ",
                {"size": 10.5},
            ),
            ("five companies writing $15–25,000 checks", {"size": 10.5, "bold": True}),
            (
                ". CSP already has the brands in the tent. The missing product is a cash sponsorship with a named outcome, not another logo on a tent.",
                {"size": 10.5},
            ),
        ],
    )
    para(
        doc,
        "What would make me change the plan: if the day-30 audit shows twenty existing $10,000+ donors, we shift from acquisition to retention and multi-year asks, and we pull corporate forward. If the board cannot or will not make introductions, corporate slips and we put more weight on DAF holders already in the file. I would rather rewrite the mix in October than defend a pretty table in May.",
    )

    heading(doc, "D. Sources consulted")
    para(
        doc,
        "citysurfproject.com (programs, partners, events, Surf-a-Thon 2026, donate, economic statement); Form 990s via ProPublica, EIN 47-2091985 (FY15–FY25); Cause IQ (FY25: $1,110,667 revenue, $1,134,574 expenses, 21 staff); Grantable / funder 990s (Endaoment, Newhall, Marin CF, Goldsmith, Parks California, Bothin, Kaiser, Miranda Lux, Olympic Club, REI CAF, Alaska Airlines); GuideStar board list (2026); SF DCYF 2024–29 RFP and Youth Surfing score report ($150,000 vs. $323,741 requested; SILT listed as funded); Coastal Commission Whale Tail (Feb 2024, $50,000); SF Standard (June 22, 2025); Development Director posting ($105–115k). The assignment brief is treated as authoritative for the 30/20/3/10/37 mix and the $1.1M / $1.4M goals — events and government often sit inside “contributions” on the 990.",
        size=9.5,
        color=MUTED,
    )

    out = Path("/workspace/docs/city-surf-project/City-Surf-Project-FY27-Fundraising-Proposal.docx")
    doc.save(out)
    print("wrote", out, out.stat().st_size)


if __name__ == "__main__":
    build()
